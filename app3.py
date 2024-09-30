import os
import urllib.request
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import time
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.shared import Inches
import re
import pickle
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LangDocument
from langchain.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain_community.chat_models import ChatOllama
from langchain.prompts import ChatPromptTemplate
from langchain.schema.runnable import RunnableMap
from docx.image.exceptions import UnrecognizedImageError
from PIL import Image  # Pillow 라이브러리

# 한글 폰트 설정
font_path = 'C:/Windows/Fonts/malgun.ttf'
font_prop = FontProperties(fname=font_path)
plt.rcParams['font.family'] = font_prop.get_name()

# 파일명에서 특수문자 제거 및 띄어쓰기 '_'로 변환하는 함수
def clean_filename(filename):
    # 특수 문자 제거 및 공백을 _로 대체
    filename = re.sub(r'[\\/*?:"<>|]', '_', filename)
    filename = re.sub(r'[\s]+', '_', filename)  # 연속된 공백을 모두 _로 대체
    filename = re.sub(r'[^\w\s]', '', filename)  # 알파벳, 숫자, 언더스코어를 제외한 모든 문자 제거
    return filename.strip('_')  # 맨 앞과 맨 뒤에 있는 언더스코어 제거


# 폴더 존재 여부 확인 및 생성
def ensure_folder_exists(folder_path):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

# 가격에서 숫자만 추출하는 함수
def extract_price(price_text):
    return int(re.sub(r'[^\d]', '', price_text))

# 중복된 제목 제거하는 함수
def remove_duplicate_titles(rows):
    seen_titles = set()
    unique_rows = []
    for row in rows:
        title = row[0]
        if title not in seen_titles:
            unique_rows.append(row)
            seen_titles.add(title)
    return unique_rows

# 이미지 다운로드 함수
def download_image(img_url, img_folder, img_name):
    ensure_folder_exists(img_folder)
    img_name_cleaned = clean_filename(img_name)  # 특수문자 제거 및 띄어쓰기 변환
    img_path = os.path.join(img_folder, f"{img_name_cleaned}.jpg")
    
    try:
        urllib.request.urlretrieve(img_url, img_path)
        if os.path.exists(img_path):
            return img_path
        else:
            print(f"이미지 다운로드 실패: {img_path}")
            return None
    except Exception as e:
        print(f"이미지 다운로드 오류: {e}")
        return None


# 이미지 형식 변환 함수
def convert_image_to_jpeg(img_path):
    try:
        if os.path.exists(img_path):  # 경로에 이미지가 있는지 확인
            with Image.open(img_path) as img:
                if img.format not in ['JPEG', 'JPG']:  # 이미지 형식이 지원되지 않으면 변환
                    img = img.convert('RGB')
                    new_img_path = img_path.rsplit('.', 1)[0] + ".jpg"
                    img.save(new_img_path, 'JPEG')
                    return new_img_path
                else:
                    return img_path  # 이미지가 이미 지원되는 형식인 경우 그대로 반환
        else:
            print(f"파일을 찾을 수 없습니다: {img_path}")
            return None
    except Exception as e:
        print(f"이미지 변환 오류: {e}")
        return None


# CSV 파일 로드 함수
def load_existing_data(folder_path, search_term):
    csv_filename = os.path.join(folder_path, f"{clean_filename(search_term)}_items.csv")
    if os.path.exists(csv_filename):
        df = pd.read_csv(csv_filename)
        if not df.empty:
            return df
    return None

# 크롤링하는 함수 (캐시된 CSV 파일이 있으면 생략)
def crawl_items(search_term, folder_path):
    csv_filename = os.path.join(folder_path, f"{clean_filename(search_term)}_items.csv")
    
    # 기존 데이터가 있으면 로드
    df = load_existing_data(folder_path, search_term)
    if df is not None:
        average_soldout_price = df['가격'].mean()
        return csv_filename, average_soldout_price, df

    # 웹 드라이버 설정 및 크롤링
    driver = webdriver.Chrome()
    driver.maximize_window()
    driver.get('https://web.joongna.com')
    time.sleep(3)

    search_box = driver.find_element(By.CLASS_NAME, 'ga4_main_top_search')
    search_box.send_keys(search_term)
    search_box.send_keys(Keys.RETURN)
    time.sleep(5)

    items = driver.find_elements(By.CSS_SELECTOR, 'li > div > a')

    rows = []
    img_folder = os.path.join(folder_path, "img")  # 이미지 폴더 경로 설정
    ensure_folder_exists(img_folder)

    for item in items:
        try:
            title = item.find_element(By.CSS_SELECTOR, 'h2').text
            price = item.find_element(By.CSS_SELECTOR, 'div.font-semibold').text
            link = item.get_attribute('href')
            img_url = item.find_element(By.CSS_SELECTOR, 'img').get_attribute('src')  # 이미지 URL
            
            # 이미지 다운로드 및 경로 저장
            img_path = download_image(img_url, img_folder, title)
            rows.append([title, extract_price(price), link, img_path])
        except Exception as e:
            print(f"오류 발생: {e}")

    driver.quit()

    # 중복된 제목 제거
    rows = remove_duplicate_titles(rows)

    # 데이터프레임 저장 및 평균 가격 계산
    df = pd.DataFrame(rows, columns=['제목', '가격', '링크', '이미지'])
    df.to_csv(csv_filename, index=False, encoding='utf-8')
    average_soldout_price = df['가격'].mean()

    return csv_filename, average_soldout_price, df

# 모델 저장 함수
def save_model(db, model_filename):
    with open(model_filename, 'wb') as model_file:
        pickle.dump(db, model_file)
    print(f"모델이 {model_filename} 파일에 저장되었습니다.")

# 모델 불러오기 함수
def load_model(model_filename):
    with open(model_filename, 'rb') as model_file:
        db = pickle.load(model_file)
    print(f"모델이 {model_filename} 파일에서 불러와졌습니다.")
    return db

# 질문에 맞춰 데이터를 기반으로 자동 응답 생성
def process_query_with_model(quest, retriever):
    context_docs = retriever.get_relevant_documents(quest)

    # 질문 유형 자동 분류 및 적절한 응답 생성
    prompt_template = """
    너는 중고거래 품목 데이터를 기반으로 정보를 대답하는 챗봇이야. 
    반드시 모든 대답은 한글로 해주세요. 
    사용자의 질문에 맞춰 관련 데이터를 제공하세요.
    추천에 대한 답변을 줄때는 중복된 [제목]을 가진 데이터를 제거하고 알려줘.
    추천을 할때는 평균 가격하고 유사한 제품만 추천해줘. 
    품목들을 테이블화 시켜서 이쁘게 보여줘.
    제품 추천할때 무엇을 기준으로 추천 했는지도 정리해서 설명해줘.
    추천할 때, 소비자에게 합리적이고 효용이 있는지 단계별로 생각하고 그 상세내용을 반드시 작성해.
    단계별로 차근차근 생각하고 알려주세요.
    

    질문: {question}
    문서: {context}
    
    """
    
    prompt = ChatPromptTemplate.from_template(prompt_template)
    llm = ChatOllama(model="gemma2:9b", temperature=0, base_url="http://127.0.0.1:11434/")
    
    # 템플릿에 맞춰 질문과 관련 문서들을 LLM에 전달
    chain = RunnableMap({
        "context": lambda x: context_docs,
        "question": lambda x: quest
    }) | prompt | llm
    
    response = chain.invoke({'question': quest}).content
    return response

# 가격 보여줄 함수
def format_number(num):
    num_str = str(num)
    if len(num_str) <= 3:
        return num_str  # 1000 미만일 경우 그냥 출력

    # 만 단위로 나누어 적당히 표현
    if len(num_str) > 3:
        만 = num_str[:-4]  # 앞의 자리수 (십만 단위)
        천 = num_str[-4:-3]  # 마지막 네 자리에서 세 번째 자리 (천 단위)
        if 천 == '0':
            return f"약 {만}만"
        return f"약 {만}만 {천}천"

# 이미지를 제목과 매칭하여 삽입하는 함수
def insert_images_from_titles(doc, folder_path, titles):
    img_folder = os.path.join(folder_path, "img")
    
    for title in titles:
        cleaned_title = clean_filename(title)  # 특수문자 처리 및 띄어쓰기 변환
        img_path = os.path.join(img_folder, f"{cleaned_title}.jpg")
        
        # 이미지 경로가 유효한지 확인
        img_path = convert_image_to_jpeg(img_path)
        
        if img_path and os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(1.25))  # 이미지 삽입
            except UnrecognizedImageError:
                doc.add_paragraph(f"이미지를 삽입할 수 없습니다: {title}")
        else:
            doc.add_paragraph(f"이미지 없음: {title}")


# 보고서 생성 함수
def generate_report(search_term, quest, chatbot_response, df, average_soldout_price, folder_path):
    doc = Document()

    # 보고서 제목
    doc.add_heading(f"{search_term} 보고서", 0)

    # 질문 추가
    doc.add_paragraph(f"질문: {quest}")

    # 평균 판매 가격
    average_str = format_number(round(average_soldout_price))
    doc.add_paragraph(f"평균 판매 가격: {average_str}원")

    # 챗봇의 답변
    doc.add_heading("챗봇의 답변", level=1)
    doc.add_paragraph(chatbot_response)

    # 추천 품목 리스트
    doc.add_heading("추천 품목 리스트", level=1)

    # 테이블 생성: 제목, 가격, 링크, 이미지
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '제목'
    hdr_cells[1].text = '가격'
    hdr_cells[2].text = '링크'
    hdr_cells[3].text = '이미지'

    # 테이블에 데이터 추가
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['제목']
        row_cells[1].text = f"{row['가격']}원"
        row_cells[2].text = row['링크']

        img_path = row['이미지']
        if img_path and os.path.exists(img_path):
            row_cells[3].text = img_path  # 이미지 경로 추가
            try:
                doc.add_picture(img_path, width=Inches(1.25))  # 이미지 삽입
            except UnrecognizedImageError:
                print(f"UnrecognizedImageError: 이미지가 인식되지 않음 - {img_path}")
                row_cells[3].text = "이미지를 삽입할 수 없습니다."  # 오류가 있을 경우 텍스트로 대체
        else:
            row_cells[3].text = "이미지 없음"

    # 챗봇 응답에서 제목 추출 및 이미지 삽입
    titles_from_chatbot = [line.split('\t')[0] for line in chatbot_response.split('\n') if line]  # 탭이나 특정 구분자로 제목 추출
    insert_images_from_titles(doc, folder_path, titles_from_chatbot)

    # 보고서 저장
    report_filename = os.path.join(folder_path, f"{clean_filename(search_term)}_보고서_{clean_filename(quest)}.docx")
    doc.save(report_filename)
    print(f"보고서가 {report_filename} 파일로 저장되었습니다.")
    return report_filename

# 메인 함수: Streamlit 웹 앱의 메인 로직
def main():
    st.title("중고거래 품목 크롤링 및 챗봇")

    search_term = st.text_input("중고거래 품목 검색어를 입력하세요")
    quest = st.text_input("중고거래 품목에 대해 질문하세요")
    update_clicked = st.sidebar.button("데이터 갱신")
    clicked_button = st.button("제출")

    if clicked_button or update_clicked:
        # 검색어에 맞는 폴더 생성
        folder_path = os.path.join(os.getcwd(), f"{clean_filename(search_term)}_데이터")
        ensure_folder_exists(folder_path)
        
        if update_clicked:
            # 데이터 갱신 버튼을 클릭하면 기존 CSV 파일 삭제
            csv_filename = os.path.join(folder_path, f"{clean_filename(search_term)}_items.csv")
            if os.path.exists(csv_filename):
                os.remove(csv_filename)

        # 크롤링 또는 기존 파일 로드
        csv_filename, average_soldout_price, df = crawl_items(search_term, folder_path)
        st.success(f"크롤링 완료. {csv_filename} 파일 생성됨.")
        average_str = format_number(round(average_soldout_price))
        st.write(f"평균 판매 가격: {average_str}원")
        st.write(df)

        # 모델 파일명 생성
        model_filename = os.path.join(folder_path, f"{clean_filename(search_term)}_model.pkl")

        # 모델이 이미 저장된 경우 로드하고, 아니면 새로 생성 후 저장
        if os.path.exists(model_filename) and not update_clicked:
            db = load_model(model_filename)
        else:
            # CSV 데이터를 문서로 변환하여 모델 학습
            documents = [LangDocument(page_content=f"{row['제목']}, {row['가격']}, {row['링크']}") for _, row in df.iterrows()]
            embedding_function = SentenceTransformerEmbeddings(model_name="jhgan/ko-sroberta-multitask")
            db = FAISS.from_documents(documents, embedding_function)
            save_model(db, model_filename)

        retriever = db.as_retriever(search_type="similarity", search_kwargs={'k': 10, 'fetch_k': 100})

        # 질문에 대한 자동 응답 생성
        chatbot_response = process_query_with_model(quest, retriever)
        st.write(chatbot_response)

        # 보고서 생성
        report_filename = generate_report(search_term, quest, chatbot_response, df, average_soldout_price, folder_path)
        st.write(f"보고서가 생성되었습니다: {report_filename}")

if __name__ == "__main__":
    main()
