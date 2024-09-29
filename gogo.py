import os
import csv
import urllib.request
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import time
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LangDocument
from langchain.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain_community.chat_models import ChatOllama
from langchain.prompts import ChatPromptTemplate
from langchain.schema.runnable import RunnableMap
import requests
from bs4 import BeautifulSoup
import pickle

# 한글 폰트 설정
plt.rcParams['font.family'] = 'Malgun Gothic'

# 파일명에서 특수문자를 제거하는 함수
def clean_filename(filename):
    return filename.replace('/', '_').replace('\\', '_').replace('*', '_')

# 가격에서 숫자만 추출하는 함수
def extract_price(price_text):
    return int(''.join(filter(str.isdigit, price_text)))

# 보고서 작성 함수 (표 포함)
def create_report(search_term, avg_price, graph_filename, chatbot_response, recommended_items):
    doc = Document()
    doc.add_heading(f'{search_term} 중고거래 시세 보고서', 0)
    doc.add_paragraph(f'평균 판매완료 가격: {avg_price} 원')

    doc.add_heading('챗봇의 의견', level=1)
    doc.add_paragraph(f"챗봇 응답: {chatbot_response}")

    doc.add_heading('추천 품목', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '제목'
    hdr_cells[1].text = '가격'
    hdr_cells[2].text = '링크'

    for index, row in recommended_items.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['제목']
        row_cells[1].text = f"{row['가격']} 원"
        row_cells[2].text = row['링크']

    if graph_filename and os.path.exists(graph_filename):
        doc.add_paragraph('시세 그래프:')
        doc.add_picture(graph_filename, width=Inches(5))

    report_path = f"{search_term}_중고거래_보고서.docx"
    doc.save(report_path)
    return report_path

# 크롤링 함수 (품목)
def crawl_items(search_term):
    csv_filename = clean_filename(f"{search_term}_items.csv")
    if os.path.exists(csv_filename):
        df = pd.read_csv(csv_filename)
        return df, df['가격'].mean()

    driver = webdriver.Chrome()
    driver.get('https://web.joongna.com')
    search_box = driver.find_element(By.CLASS_NAME, 'ga4_main_top_search')
    search_box.send_keys(search_term)
    search_box.send_keys(Keys.RETURN)
    time.sleep(5)

    items = driver.find_elements(By.CSS_SELECTOR, 'li > div > a')
    data = []
    for item in items:
        try:
            title = item.find_element(By.CSS_SELECTOR, 'h2').text
            price = item.find_element(By.CSS_SELECTOR, 'div.font-semibold').text
            link = item.get_attribute('href')
            data.append([title, price, link])
        except:
            continue
    driver.quit()

    df = pd.DataFrame(data, columns=['제목', '가격', '링크'])
    df['가격'] = df['가격'].apply(extract_price)
    df.to_csv(csv_filename, index=False)
    return df, df['가격'].mean()

# 중고 거래 팁 크롤링 함수
def scrape_tips():
    url = "https://namu.wiki/w/%EC%A4%91%EA%B3%A0%EA%B1%B0%EB%9E%98"
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    tips = []
    titles = soup.find_all(class_='IL0R8WAn')
    contents = soup.find_all(class_='flUNFbH-')

    for title, content in zip(titles, contents):
        tips.append({
            'title': title.get_text().strip(),
            'content': content.get_text().strip()
        })
    
    return tips

# 모델 저장 및 불러오기
def save_model(db, model_filename):
    with open(model_filename, 'wb') as file:
        pickle.dump(db, file)

def load_model(model_filename):
    with open(model_filename, 'rb') as file:
        return pickle.load(file)

# 추천 질문 항목 표시
def show_recommended_questions():
    st.sidebar.markdown("### 추천 질문 항목")
    st.sidebar.markdown("- 어떤 가격에 구매하면 좋을까?")
    st.sidebar.markdown("- 제일 싼 가격의 물품의 가격과 링크를 알려줘")
    st.sidebar.markdown("- 중고 거래 시 중요한 팁은?")

# 메인 함수
def main():
    st.title("중고거래 품목 크롤링 및 챗봇")

    # 검색어 입력 및 크롤링
    search_term = st.text_input("중고거래 품목 검색어를 입력하세요")
    quest = st.text_input("중고거래 품목에 대해 질문하세요 (예: 어떤 가격에 구매하면 좋을까?)")
    update_clicked = st.sidebar.button("데이터 갱신")
    clicked_button = st.button("제출")

    recommended_items = pd.DataFrame()

    if clicked_button or update_clicked:
        if update_clicked:
            csv_filename = clean_filename(f"{search_term}_items.csv")
            if os.path.exists(csv_filename):
                os.remove(csv_filename)

        df, avg_price = crawl_items(search_term)
        st.write(f"평균 가격: {avg_price} 원")
        st.write(df)

        model_filename = clean_filename(f"{search_term}_model.pkl")
        if os.path.exists(model_filename) and not update_clicked:
            db = load_model(model_filename)
        else:
            documents = [LangDocument(page_content=", ".join([f"{col}: {row[col]}" for col in df.columns])) for index, row in df.iterrows()]
            embedding_function = SentenceTransformerEmbeddings(model_name="jhgan/ko-sroberta-multitask")
            db = FAISS.from_documents(documents, embedding_function)
            save_model(db, model_filename)

        retriever = db.as_retriever(search_type="similarity", search_kwargs={'k': 10})

        if "어떤 가격에 구매하면 좋을까?" in quest:
            chatbot_response = f"평균 판매 가격은 {avg_price} 원입니다. 이 가격대에 맞는 상품을 추천해드리겠습니다."
            recommended_items = df[(df['가격'] > avg_price * 0.9) & (df['가격'] < avg_price * 1.1)]
        elif "제일 싼 가격의 물품" in quest:
            cheapest_item = df.loc[df['가격'].idxmin()]
            chatbot_response = f"제일 싼 물품은 {cheapest_item['제목']}입니다. 가격은 {cheapest_item['가격']} 원입니다."
        elif "중고 거래 시 중요한 팁" in quest:
            tips = scrape_tips()
            chatbot_response = "중요한 중고 거래 팁입니다:\n" + "\n".join([f"{tip['title']} - {tip['content']}" for tip in tips])
        else:
            prompt = ChatPromptTemplate.from_template("""
            너는 중고거래 품목 데이터를 기반으로 정보를 대답하는 챗봇이야. 모든 대답은 한글로 작성해주세요.
            """)
            llm = ChatOllama(model="gemma2:9b")
            chain = RunnableMap({
                "context": lambda x: retriever.get_relevant_documents(x['question']),
                "question": lambda x: x['question']
            }) | prompt | llm
            chatbot_response = chain.invoke({'question': quest}).content

        st.write(chatbot_response)

        report_path = create_report(search_term, avg_price, None, chatbot_response, recommended_items)
        st.success(f"보고서가 생성되었습니다: {report_path}")

    show_recommended_questions()

if __name__ == "__main__":
    main()
